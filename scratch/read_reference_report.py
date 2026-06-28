
import sys
import subprocess
import os

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    from docx import Document
except ImportError:
    install("python-docx")
    from docx import Document

def read_docx(file_path):
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
                
    return "\n".join(full_text)

if __name__ == "__main__":
    # Trying common possible names
    paths = [
        r"C:\Users\Admin\Downloads\Arthi_project\Phase_2_report.docx",
        r"C:\Users\Admin\Downloads\Arthi_project\reportformat.docx",
        r"C:\Users\Admin\Downloads\Arthi_project check report2.docx"
    ]
    
    for path in paths:
        print(f"--- CONTENT OF {path} ---")
        text = read_docx(path)
        if text:
            print(text[:2000]) # Print first 2000 chars
            print("--- END ---")
        else:
            print("Could not read or file doesn't exist.")
