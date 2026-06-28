import subprocess
import sys
import os

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx...")
    install("python-docx")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('start', 'top', 'bottom', 'end', 'left', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def create_report():
    doc = Document()

    # --- FRONT PAGE ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Sanjeevan Group of Institutions, Panhala\n')
    run.bold = True
    run.font.size = Pt(18)
    run = p.add_run('Sanjeevan Knowledge City, Somwar Peth, Panhala, Kolhapur - 416201\n')
    run.font.size = Pt(12)
    doc.add_paragraph('\n\n')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Department of Computer Science and Engineering\n')
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph('\n')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PROJECT PHASE-II REPORT ON\n')
    run.font.size = Pt(14)
    run = p.add_run('"MULTILINGUAL TRANSLATION: OFFLINE MULTILINGUAL AI TRANSLATION SYSTEM"')
    run.bold = True
    run.font.size = Pt(20)
    doc.add_paragraph('\n\n')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Submitted in partial fulfillment of the requirements for the award of degree of\n')
    run.font.size = Pt(12)
    run = p.add_run('Bachelor of Technology (B. Tech)')
    run.bold = True
    doc.add_paragraph('\n')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SUBMITTED BY\n')
    run.font.size = Pt(12)
    run = p.add_run('Mr. Arman Mulla\n')
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph('\n')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('UNDER THE GUIDANCE OF\n')
    run.font.size = Pt(12)
    run = p.add_run('Prof. R. S. Nejkar')
    run.bold = True
    run.font.size = Pt(14)
    doc.add_page_break()

    # --- CERTIFICATE ---
    doc.add_heading('CERTIFICATE', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n')
    cert_text = (
        "This is to certify that Project entitled \"MULTILINGUAL TRANSLATION: OFFLINE MULTILINGUAL AI TRANSLATION SYSTEM\" "
        "is a bonafied Project work carried out by Mr. Arman Mulla in partial fulfillment of the requirements "
        "for the award of Bachelor of Technology (B. Tech) in Computer Science and Engineering by "
        "Dr. Babasaheb Ambedkar Technological University, Lonere for the academic year 2025-2026."
    )
    doc.add_paragraph(cert_text)
    doc.add_paragraph('\n\n\n\n')
    table = doc.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "Guide\nProf. R. S. Nejkar"
    table.cell(0, 1).text = "HOD\nProf. R. S. Nejkar"
    table.cell(0, 2).text = "Principal\nDr. S. N. Jain"
    doc.add_page_break()

    # --- CHAPTER 1: INTRODUCTION ---
    doc.add_heading('CHAPTER 1: INTRODUCTION', level=1)
    doc.add_heading('1.1 Background', level=2)
    doc.add_paragraph("Multilingual Translation is an offline translation application providing text, audio, video, and PDF translation securely.")
    doc.add_heading('1.2 Problem Definition', level=2)
    doc.add_paragraph("Current systems need internet and risk privacy. Multilingual Translation offers a secure offline alternative.")
    doc.add_heading('1.3 Objectives', level=2)
    doc.add_paragraph("To build a 100% offline system; Implement Whisper and M2M100 models; Support multiple file formats.", style='List Bullet')
    doc.add_page_break()

    # --- CHAPTER 3: SYSTEM DESIGN ---
    doc.add_heading('CHAPTER 3: SYSTEM DESIGN', level=1)
    doc.add_heading('3.1 System Architecture', level=2)
    
    # SIMPLE BOXED ARCHITECTURE (Matching user image style)
    def add_simple_layer(title, items):
        # Add Title Paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        
        # Add Table for Items (Boxes side by side)
        table = doc.add_table(rows=1, cols=len(items))
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, item in enumerate(items):
            cell = table.cell(0, i)
            cell.text = item
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_border(cell, top={"sz": 6}, bottom={"sz": 6}, start={"sz": 6}, end={"sz": 6})
        
        # Add Arrow
        arrow_p = doc.add_paragraph("↓")
        arrow_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_simple_layer("1. CLIENT / UI LAYER", ["HTML/Tailwind", "Dashboard", "File Upload", "History Viewer"])
    add_simple_layer("2. FASTAPI BACKEND LAYER", ["API Routes", "Handlers", "UUID Gen", "Static Files"])
    add_simple_layer("3. AI ENGINE LAYER (LOCAL)", ["Whisper (STT)", "M2M100 (MT)", "MoviePy", "PyPDF2"])
    add_simple_layer("4. DATABASE LAYER (SQLITE)", ["History", "Translations", "Language Logs"])
    
    # Final layer without arrow
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("5. STORAGE LAYER (LOCAL)")
    run.bold = True
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    items = ["Model Weights", "Temp Uploads"]
    for i, item in enumerate(items):
        cell = table.cell(0, i)
        cell.text = item
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border(cell, top={"sz": 6}, bottom={"sz": 6}, start={"sz": 6}, end={"sz": 6})

    doc.add_page_break()

    # --- CHAPTER 4: IMPLEMENTATION ---
    doc.add_heading('CHAPTER 4: IMPLEMENTATION', level=1)
    doc.add_heading('4.1 DFD Level 0 (Context Diagram)', level=2)
    doc.add_paragraph("Input: User Uploads (Audio/Video/PDF/Text) -> PROCESS: Multilingual Translation Engine -> Output: Translated Data.")
    
    doc.add_heading('4.2 DFD Level 1 (Data Flow)', level=2)
    doc.add_paragraph("1. Frontend sends file to FastAPI.\n2. Backend extracts media using MoviePy.\n3. AI Models (Whisper/M2M100) process data.\n4. Database saves result and sends to UI.")

    doc.add_page_break()

    # --- CHAPTER 5: CONCLUSION ---
    doc.add_heading('CHAPTER 5: CONCLUSION', level=1)
    doc.add_paragraph("Multilingual Translation provides a secure, offline, and high-performance solution for multilingual translation.")

    filename = "Multilingual_Translation_Official_Project_Report.docx"
    doc.save(filename)
    print(f"Simple Report generated: {filename}")

if __name__ == "__main__":
    create_report()
