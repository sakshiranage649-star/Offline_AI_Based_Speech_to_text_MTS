import subprocess
import sys
import os

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...")
    install("python-docx")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_setup_guide():
    doc = Document()

    # --- Title ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Multilingual Translation: Offline Multilingual AI Translation System\n')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Complete Project Setup & Installation Guide')
    run.bold = True
    run.font.size = Pt(18)
    
    doc.add_paragraph('\n')

    # --- Section 1: Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "Multilingual Translation is an AI-powered translation application designed to work completely offline. "
        "It supports text translation, speech-to-text (STT), and document translation (PDF/Video/Audio). "
        "This guide will walk you through the process of setting up the environment and running the application on your PC."
    )

    # --- Section 2: Prerequisites ---
    doc.add_heading('2. System Prerequisites', level=1)
    doc.add_paragraph("Before starting the setup, ensure your system meets the following requirements:", style='Body Text')
    
    prereqs = [
        "Operating System: Windows 10 or 11 (64-bit recommended)",
        "Python: Version 3.8 to 3.11 (Download from python.org)",
        "RAM: Minimum 8GB (16GB recommended for faster AI processing)",
        "Disk Space: At least 10GB free (for AI models and dependencies)",
        "Internet: Required ONLY for the initial setup/installation phase."
    ]
    for item in prereqs:
        doc.add_paragraph(item, style='List Bullet')

    # --- Section 3: Installation Steps ---
    doc.add_heading('3. Installation Steps', level=1)

    doc.add_heading('Step 1: Extract Project Files', level=2)
    doc.add_paragraph("Extract the 'Multilingual Translation' project folder to a convenient location on your PC (e.g., Desktop or C: drive).")

    doc.add_heading('Step 2: Run the Installer', level=2)
    doc.add_paragraph(
        "To automate the installation of dependencies and download AI models, double-click on the 'install_models.bat' file. "
        "This script will:"
    )
    doc.add_paragraph("Install all required Python libraries (FastAPI, Transformers, Torch, etc.)", style='List Number')
    doc.add_paragraph("Download the Whisper (STT) and M2M100 (Translation) models from HuggingFace.", style='List Number')
    doc.add_paragraph("Verify the models are working correctly.", style='List Number')
    
    doc.add_paragraph(
        "NOTE: The model download is approximately 5GB. Depending on your internet speed, this may take 10-30 minutes.",
        style='Intense Quote'
    )

    doc.add_heading('Step 3: Database Setup', level=2)
    doc.add_paragraph(
        "The application uses a local SQLite database ('transly.db'). "
        "The backend is configured to automatically create and initialize the database on the first run. "
        "No manual SQL installation is required."
    )

    # --- Section 4: Running the Application ---
    doc.add_heading('4. Running the Application', level=1)
    doc.add_paragraph(
        "Once the installation is complete, follow these steps to start the application:"
    )
    doc.add_paragraph("Double-click on the 'run.bat' file in the project root folder.", style='List Number')
    doc.add_paragraph("A command prompt will open, starting the FastAPI server.", style='List Number')
    doc.add_paragraph("Wait for the message 'Started server process' to appear.", style='List Number')
    doc.add_paragraph("The application will automatically open in your default web browser at http://127.0.0.1:8000", style='List Number')

    # --- Section 5: Usage and Credentials ---
    doc.add_heading('5. Default Credentials', level=1)
    doc.add_paragraph("You can log in to the application using the following default accounts:")
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Role'
    hdr_cells[1].text = 'Username'
    hdr_cells[2].text = 'Password'
    
    row_user = table.rows[1].cells
    row_user[0].text = 'Admin'
    row_user[1].text = 'Admin'
    row_user[2].text = 'Admin123'
    
    row_guest = table.rows[2].cells
    row_guest[0].text = 'User'
    row_guest[1].text = 'Any (Create via Register)'
    row_guest[2].text = 'Any'

    # --- Section 6: Troubleshooting ---
    doc.add_heading('6. Troubleshooting', level=1)
    
    troubles = [
        "Python Not Found: Ensure Python is added to your system PATH during installation.",
        "Port 8000 Busy: If another app is using port 8000, the server won't start. Close other Python servers or change the port in backend/app.py.",
        "Model Error: If you get a 'Model not found' error, run 'install_models.bat' again to ensure downloads completed.",
        "Slow Translation: The first translation after starting the server might be slow as models are loaded into RAM."
    ]
    for item in troubles:
        doc.add_paragraph(item, style='List Bullet')

    # --- Footer ---
    doc.add_paragraph('\n\n')
    p = doc.add_paragraph('Developed by Arman Mulla')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('\nMultilingual Translation Project © 2026')
    run.font.size = Pt(10)

    filename = "Multilingual_Translation_Setup_Guide.docx"
    doc.save(filename)
    print(f"Setup Guide generated successfully: {filename}")

if __name__ == "__main__":
    create_setup_guide()
