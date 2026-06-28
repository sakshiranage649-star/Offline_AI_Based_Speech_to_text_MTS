import subprocess
import sys
import os

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...")
    install("python-docx")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_viva_guide():
    doc = Document()

    # --- Title ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Multilingual Translation: Viva & Project Examination Guide\n')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Expert Answers to Potential Examiner Questions')
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph('\n')

    questions = [
        {
            "q": "1. What is the core objective of the 'Multilingual Translation' project?",
            "a": "The core objective is to provide a privacy-focused, 100% offline multilingual translation system. It allows users to translate text, documents, audio, and video without needing an internet connection, ensuring data security and accessibility in remote areas."
        },
        {
            "q": "2. Which specific AI models are used in this project?",
            "a": "We use two primary models: 'facebook/m2m100_418M' for high-quality multilingual translation (supporting 100 languages) and 'openai/whisper-small' for accurate Speech-to-Text (STT) transcription."
        },
        {
            "q": "3. Why did you choose SQLite as the database engine?",
            "a": "SQLite was chosen because it is serverless and zero-configuration. Since Multilingual Translation is designed for offline local use, SQLite allows the entire application to be portable (contained in one folder) without requiring the user to install a database server like MySQL."
        },
        {
            "q": "4. Explain the backend architecture of the project.",
            "a": "The project uses FastAPI, a modern, high-performance web framework for Python. It handles asynchronous requests, serves the frontend files, and provides API endpoints for translation, user management, and report generation."
        },
        {
            "q": "5. How is user security and password management handled?",
            "a": "User passwords are never stored in plain text. We use SHA-256 hashing to secure passwords in the database. When a user logs in, their input is hashed and compared with the stored hash to verify identity."
        },
        {
            "q": "6. What is the purpose of the 'Repetition Penalty' in your translation logic?",
            "a": "In AI translation, models sometimes get stuck in a loop, repeating the same word. We implemented a 'repetition_penalty' of 1.4 and 'no_repeat_ngram_size=2' to mathematically prevent these loops and ensure natural-sounding translations."
        },
        {
            "q": "7. How does the application process video files for translation?",
            "a": "The process involves three steps: 1. Extracting audio from the video using the MoviePy library. 2. Transcribing that audio into text using the Whisper model. 3. Translating the extracted text into the target language using M2M100."
        },
        {
            "q": "8. How does the 'Token System' work for users?",
            "a": "Each user is assigned a specific number of tokens (e.g., 1000). Every translation consumes tokens based on the character count of the text. This simulates a real-world SaaS billing model and allows admins to manage user limits."
        },
        {
            "q": "9. How does the system handle hardware variations, specifically GPU vs. CPU?",
            "a": "The application detects if a CUDA-enabled NVIDIA GPU is available using PyTorch. If found, it loads models onto the GPU for 10x faster processing. If not, it automatically falls back to the CPU, ensuring it works on all computers."
        },
        {
            "q": "10. What challenges did you face with document translation (PDFs)?",
            "a": "Extracting clean text from PDFs can be difficult due to varying formats and encoding. We used the PyPDF2 library to extract text while maintaining as much structural integrity as possible before sending it to the translation engine."
        },
        {
            "q": "11. Explain the role of the '.env' file in this project.",
            "a": "The .env file stores configuration variables like DB_TYPE and DB_NAME. This allows us to change database settings (e.g., switching from SQLite to MySQL) without modifying the actual source code."
        },
        {
            "q": "12. What are the key capabilities of the Admin Dashboard?",
            "a": "The Admin Dashboard provides real-time reports on total users, translation traffic (last 7 days), token leaders, and language popularity. Admins can also update user tokens and manage global application settings."
        },
        {
            "q": "13. How did you ensure the frontend is responsive and user-friendly?",
            "a": "We used semantic HTML5 and Vanilla CSS with a modern 'Glassmorphism' design. It features smooth transitions, responsive layouts for different screen sizes, and dynamic error handling to improve the user experience."
        },
        {
            "q": "14. What are the advantages of using M2M100 over the Google Translate API?",
            "a": "M2M100 is an 'Open Source' and 'Local' model. Unlike Google Translate API, it requires no internet, has no per-request cost, and guarantees 100% data privacy since no data ever leaves the local machine."
        },
        {
            "q": "15. How can this project be scaled or improved in the future?",
            "a": "Future improvements include: 1. Adding OCR (Optical Character Recognition) to translate text from images. 2. Implementing a real-time voice translation mode. 3. Containerizing the app with Docker for easier cloud deployment."
        },
        {
            "q": "16. Why did you use FastAPI instead of Flask or Django?",
            "a": "FastAPI is significantly faster due to its asynchronous nature. It also provides automatic Swagger documentation and built-in data validation, which makes it easier to build and maintain robust APIs compared to Flask."
        },
        {
            "q": "17. Explain the 'translations' table schema in the database.",
            "a": "The table includes: id (primary key), user_id (foreign key), source_lang, target_lang, source_text, translated_text, and a timestamp. This allows us to track history and generate usage analytics for the dashboard."
        }
    ]

    for item in questions:
        q_para = doc.add_paragraph()
        run_q = q_para.add_run(item["q"])
        run_q.bold = True
        run_q.font.size = Pt(12)
        
        a_para = doc.add_paragraph()
        run_a = a_para.add_run("Answer: ")
        run_a.bold = True
        a_para.add_run(item["a"])
        doc.add_paragraph() # Spacer

    # --- Footer ---
    doc.add_paragraph('\n\n')
    p = doc.add_paragraph('Prepared for Multilingual Translation Project Viva')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('\nGenerated by AI Assistant')
    run.font.size = Pt(10)

    filename = "Multilingual_Translation_Viva_Questions_Expert_Guide.docx"
    doc.save(filename)
    print(f"Viva Guide generated successfully: {filename}")

if __name__ == "__main__":
    create_viva_guide()
