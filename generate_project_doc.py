import os
import sys
import subprocess

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...")
    install("python-docx")
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_project_info_doc():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Multilingual Translation: Project Information & Technical Documentation', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("This document provides a comprehensive overview of the 'Multilingual Translation' system, its architecture, and the technologies used in its development.")

    # --- Section 1: Overview ---
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph(
        "Multilingual Translation is an offline-first multilingual translation application. It enables users to translate text, documents (PDF), "
        "and media files (Audio/Video) without an internet connection, ensuring 100% data privacy and security."
    )

    # --- Section 2: Technical Stack ---
    doc.add_heading('2. Technical Stack (Every Single Thing Used)', level=1)
    
    doc.add_heading('Frontend (The User Interface)', level=2)
    doc.add_paragraph("• HTML5 & CSS3: Core structure and premium styling.", style='List Bullet')
    doc.add_paragraph("• Vanilla JavaScript: API integration and dynamic UI updates.", style='List Bullet')
    doc.add_paragraph("• Lucide Icons: Modern iconography.", style='List Bullet')
    doc.add_paragraph("• Chart.js: Administrative data visualizations.", style='List Bullet')

    doc.add_heading('Backend (The Engine)', level=2)
    doc.add_paragraph("• Python 3.x: Main development language.", style='List Bullet')
    doc.add_paragraph("• FastAPI: High-performance web framework.", style='List Bullet')
    doc.add_paragraph("• Uvicorn: ASGI server for running the app.", style='List Bullet')

    doc.add_heading('Database (Storage)', level=2)
    doc.add_paragraph("• SQLite: Local, serverless database (transly.db).", style='List Bullet')
    doc.add_paragraph("• sqlite3: Python connector for DB operations.", style='List Bullet')

    doc.add_heading('AI & Machine Learning', level=2)
    doc.add_paragraph("• Transformers (Hugging Face): Model management.", style='List Bullet')
    doc.add_paragraph("• M2M100 (418M): Multi-language translation model.", style='List Bullet')
    doc.add_paragraph("• OpenAI Whisper: High-accuracy Speech-to-Text.", style='List Bullet')
    doc.add_paragraph("• PyTorch: Deep learning framework execution.", style='List Bullet')

    doc.add_heading('Libraries & Utilities', level=2)
    doc.add_paragraph("• MoviePy: Video and audio extraction.", style='List Bullet')
    doc.add_paragraph("• PyPDF2: PDF text extraction.", style='List Bullet')
    doc.add_paragraph("• Librosa: Audio processing and resampling.", style='List Bullet')
    doc.add_paragraph("• python-multipart: Handling file uploads.", style='List Bullet')

    # --- Section 3: Project Structure ---
    doc.add_heading('3. Project Structure', level=1)
    doc.add_paragraph("• backend/: Contains app.py, database logic, and AI handlers.", style='List Bullet')
    doc.add_paragraph("• frontend/: Contains the HTML/JS/CSS client-side code.", style='List Bullet')
    doc.add_paragraph("• transly.db: The single database file storing all app data.", style='List Bullet')
    doc.add_paragraph("• requirements.txt: List of all 14+ dependencies.", style='List Bullet')
    doc.add_paragraph("• run.bat: The automated startup script.", style='List Bullet')

    # --- Section 4: Key Features ---
    doc.add_heading('4. Core Features', level=1)
    doc.add_paragraph("1. Multilingual Text Translation: Support for 100+ languages.", style='List Number')
    doc.add_paragraph("2. PDF Translation: Automated text extraction and translation.", style='List Number')
    doc.add_paragraph("3. Audio/Video Transcription: Converting speech to text locally.", style='List Number')
    doc.add_paragraph("4. User Token System: Controlling usage limits for users.", style='List Number')
    doc.add_paragraph("5. Admin Dashboard: Managing users and viewing usage reports.", style='List Number')

    # --- Section 5: Startup ---
    doc.add_heading('5. How to Run', level=1)
    doc.add_paragraph("Simply run the 'run.bat' file. It handles environment setup, dependency installation, and server startup automatically.")

    filename = "Multilingual_Translation_Project_Information.docx"
    doc.save(filename)
    print(f"Document created successfully: {filename}")

if __name__ == "__main__":
    create_project_info_doc()
